# utils.py

import markdown
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO

def export_to_word(markdown_content, references):
    """
    Convierte el contenido Markdown a un documento de Word, incluyendo las referencias.
    """
    # Convertir Markdown a HTML
    html = markdown.markdown(markdown_content)
    
    # Parsear HTML con BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')
    
    # Crear un documento de Word
    doc = Document()
    
    # Iterar sobre los elementos del HTML y agregarlos al documento
    for element in soup.descendants:
        if isinstance(element, str):
            continue  # Ignorar cadenas de texto directas
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
    
    # Agregar sección de Referencias si existen
    if references:
        doc.add_heading("Referencias", level=2)
        for ref in references:
            doc.add_paragraph(ref, style='List Number')
    
    # Guardar el documento en un objeto BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
