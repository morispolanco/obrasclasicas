import streamlit as st
import requests
import markdown
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO

# Configuración de la página
st.set_page_config(
    page_title="Generador de Estudios de Obras Clásicas",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Título de la aplicación
st.title("Generador de Estudios de Obras Clásicas")

# Descripción de la aplicación
st.markdown("""
Esta aplicación genera automáticamente estudios detallados de grandes obras clásicas en campos como la literatura, filosofía, política, entre otros.
Puedes generar secciones de análisis una por una, lo que te permite pausar el proceso en cualquier momento. Además, puedes editar la información inicial y regenerar secciones específicas si lo deseas.
Finalmente, puedes exportar el estudio a un archivo en formato Word, ya sea completo o parcial, con referencias académicas incluidas.
""")

# Inicialización de variables en el estado de la sesión
if 'title' not in st.session_state:
    st.session_state.title = ""
if 'description' not in st.session_state:
    st.session_state.description = ""
if 'table_of_contents' not in st.session_state:
    st.session_state.table_of_contents = []
if 'sections' not in st.session_state:
    st.session_state.sections = []  # Lista de diccionarios: [{'number': 1, 'title': 'Título', 'content': 'Contenido'}, ...]
if 'current_section' not in st.session_state:
    st.session_state.current_section = 1
if 'total_sections' not in st.session_state:
    st.session_state.total_sections = 10  # Total de secciones, ajustable según necesidades
if 'markdown_content' not in st.session_state:
    st.session_state.markdown_content = ""
if 'generation_complete' not in st.session_state:
    st.session_state.generation_complete = False
if 'selected_section' not in st.session_state:
    st.session_state.selected_section = None
if 'references' not in st.session_state:
    st.session_state.references = []  # Lista de referencias académicas
if 'work_type' not in st.session_state:
    st.session_state.work_type = ""  # Tipo de obra: literaria, filosófica, política, etc.

# Función para reiniciar el estado de la sesión
def reset_session():
    st.session_state.title = ""
    st.session_state.description = ""
    st.session_state.table_of_contents = []
    st.session_state.sections = []
    st.session_state.current_section = 1
    st.session_state.markdown_content = ""
    st.session_state.generation_complete = False
    st.session_state.selected_section = None
    st.session_state.references = []
    st.session_state.work_type = ""

# Función para llamar a la API de OpenRouter
def call_openrouter_api(messages, model="qwen/qwen-2.5-72b-instruct"):
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}"
    }
    data = {
        "model": model,
        "messages": messages
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content'].strip()
    except requests.exceptions.HTTPError as err:
        st.error(f"Error en la API: {err}")
        return None
    except Exception as e:
        st.error(f"Error inesperado: {e}")
        return None

# Función para generar título y descripción
def generate_title_description(work_title, work_type):
    prompt = f"""Genera un título y una descripción para un estudio detallado de la siguiente obra clásica. El estudio debe ser educativo, claro y profesional.
    
Título de la obra: {work_title}
Tipo de obra: {work_type}
    
Formato de respuesta:
Título del Estudio: [Título del estudio]
Descripción: [Descripción del estudio]
"""
    messages = [
        {"role": "user", "content": prompt}
    ]
    response = call_openrouter_api(messages)
    if response:
        # Separar el título y la descripción
        lines = response.split('\n')
        title = ""
        description = ""
        for line in lines:
            if line.startswith("Título del Estudio:"):
                title = line.replace("Título del Estudio:", "").strip()
            elif line.startswith("Descripción:"):
                description = line.replace("Descripción:", "").strip()
        return title, description
    return None, None

# Función para generar la tabla de contenidos
def generate_table_of_contents(work_title, work_type, total_sections):
    prompt = f"""Genera una tabla de contenidos para un estudio detallado de la siguiente obra clásica. La tabla debe contener {total_sections} secciones con títulos descriptivos y relevantes, adaptados al tipo de obra.
    
Título de la obra: {work_title}
Tipo de obra: {work_type}
    
Formato de respuesta:
Sección 1: [Título de la Sección 1]
Sección 2: [Título de la Sección 2]
...
Sección {total_sections}: [Título de la Sección {total_sections}]
"""
    messages = [
        {"role": "user", "content": prompt}
    ]
    response = call_openrouter_api(messages)
    if response:
        table = []
        for line in response.split('\n'):
            if line.startswith("Sección"):
                parts = line.split(":")
                if len(parts) >= 2:
                    sec_num = int(parts[0].split(" ")[1])
                    sec_title = parts[1].strip()
                    table.append({"number": sec_num, "title": sec_title, "content": "", "references": []})
        return table
    return None

# Función para generar un título de sección
def generate_section_title(work_title, work_type, section_num):
    prompt = f"""Genera un título único y descriptivo para la sección {section_num} de un estudio detallado de la siguiente obra clásica.
    
Título de la obra: {work_title}
Tipo de obra: {work_type}
    
Formato de respuesta:
Título de la Sección {section_num}: [Título Único]
"""
    messages = [
        {"role": "user", "content": prompt}
    ]
    response = call_openrouter_api(messages)
    if response:
        title = ""
        for line in response.split('\n'):
            if line.startswith(f"Título de la Sección {section_num}:"):
                title = line.replace(f"Título de la Sección {section_num}:", "").strip()
                break
        return title
    return None

# Función para generar una sección con análisis detallado
def generate_section(work_title, work_type, section_num):
    # Definir el tipo de análisis según el tipo de obra
    if work_type.lower() == "literaria":
        analysis_prompt = f"""Escribe el contenido de la sección {section_num} para un estudio detallado de la obra "{work_title}".
        
Tipo de obra: Literatura
        
La sección debe incluir análisis de personajes, técnicas narrativas, contexto histórico, biografía del autor y cualquier otro análisis relevante.
        
El contenido debe ser claro, educativo y bien estructurado, con aproximadamente 3000 tokens. Incluye referencias académicas pertinentes al final de la sección en formato APA.
"""
    elif work_type.lower() == "filosófica":
        analysis_prompt = f"""Escribe el contenido de la sección {section_num} para un estudio detallado de la obra "{work_title}".
        
Tipo de obra: Filosofía
        
La sección debe incluir estudio del autor, contexto histórico, ideas principales, discusiones académicas en torno a las ideas presentadas y cualquier otro análisis relevante.
        
El contenido debe ser claro, educativo y bien estructurado, con aproximadamente 3000 tokens. Incluye referencias académicas pertinentes al final de la sección en formato APA.
"""
    elif work_type.lower() == "política":
        analysis_prompt = f"""Escribe el contenido de la sección {section_num} para un estudio detallado de la obra "{work_title}".
        
Tipo de obra: Política
        
La sección debe incluir estudio del autor, contexto histórico, teorías políticas presentadas, impacto en la sociedad, discusiones académicas y cualquier otro análisis relevante.
        
El contenido debe ser claro, educativo y bien estructurado, con aproximadamente 3000 tokens. Incluye referencias académicas pertinentes al final de la sección en formato APA.
"""
    else:
        analysis_prompt = f"""Escribe el contenido de la sección {section_num} para un estudio detallado de la obra "{work_title}".
        
Tipo de obra: {work_type}
        
La sección debe incluir análisis detallado relevante al tipo de obra, estructurado de manera clara y educativa, con aproximadamente 3000 tokens. Incluye referencias académicas pertinentes al final de la sección en formato APA.
"""
    
    messages = [
        {"role": "user", "content": analysis_prompt}
    ]
    response = call_openrouter_api(messages)
    return response

# Función para extraer referencias del contenido generado
def extract_references(section_content):
    # Asumiendo que las referencias están al final del contenido después de una sección llamada "Referencias"
    references = []
    if "Referencias" in section_content:
        ref_section = section_content.split("Referencias")[-1]
        for line in ref_section.split('\n'):
            line = line.strip()
            if line and (line.lower().startswith("apa") is False):  # Excluir posibles indicaciones de formato
                references.append(line)
    return references

# Función para exportar a Word
def export_to_word(markdown_content, references):
    # Convertir Markdown a HTML
    html = markdown.markdown(markdown_content)
    
    # Parsear HTML con BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')
    
    # Crear un documento de Word
    doc = Document()
    
    # Iterar sobre los elementos del HTML y agregarlos al documento
    for element in soup.descendants:
        if isinstance(element, str):
            continue  # Ignorar cadenas de texto directas
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
    
    # Agregar sección de Referencias si existen
    if references:
        doc.add_heading("Referencias", level=2)
        for ref in references:
            doc.add_paragraph(ref, style='List Number')
    
    # Guardar el documento en un objeto BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Barra Lateral ---
with st.sidebar:
    st.header("Menú de Navegación")
    
    # Botón para reiniciar la generación
    if st.button("Reiniciar Generación"):
        reset_session()
        st.success("Estado de la generación reiniciado.")
    
    st.markdown("---")
    
    # Menú desplegable para ver secciones generadas
    if st.session_state.sections:
        selected_section = st.selectbox(
            "Selecciona una sección para ver:",
            [f"Sección {sec['number']}" for sec in st.session_state.sections],
            key="select_section"
        )
        section_index = int(selected_section.split(" ")[1]) - 1
        # Verificar que el índice esté dentro del rango
        if 0 <= section_index < len(st.session_state.sections):
            st.session_state.selected_section = section_index
        else:
            st.error("Sección seleccionada no válida.")
    else:
        st.info("No hay secciones generadas aún.")

# --- Sección Principal ---

# Entrada de usuario para el título de la obra
work_title = st.text_input("Ingresa el título de la obra clásica:", "")

# Selección del tipo de obra
work_type = st.selectbox(
    "Selecciona el tipo de obra:",
    ["Literaria", "Filosófica", "Política", "Otro"],
    index=0
)

# Botón para generar título, descripción y tabla de contenidos
if not st.session_state.title:
    if st.button("Generar Título, Descripción y Tabla de Contenidos"):
        if not work_title:
            st.warning("Por favor, ingresa el título de la obra para generar el estudio.")
        else:
            with st.spinner("Generando título, descripción y tabla de contenidos..."):
                title, description = generate_title_description(work_title, work_type)
                if title and description:
                    table_of_contents = generate_table_of_contents(work_title, work_type, st.session_state.total_sections)
                    if table_of_contents:
                        st.session_state.title = title
                        st.session_state.description = description
                        st.session_state.table_of_contents = table_of_contents
                        # Construir el contenido Markdown
                        st.session_state.markdown_content = f"# {title}\n\n{description}\n\n## Tabla de Contenidos\n\n"
                        for sec in table_of_contents:
                            st.session_state.markdown_content += f"{sec['number']}. {sec['title']}\n"
                        st.session_state.markdown_content += "\n"
                        # Inicializar la lista de secciones sin contenido
                        st.session_state.sections = [{"number": sec['number'], "title": sec['title'], "content": "", "references": []} for sec in table_of_contents]
                        st.success("Título, descripción y tabla de contenidos generados exitosamente.")
                        st.subheader("Título")
                        st.write(title)
                        st.subheader("Descripción")
                        st.write(description)
                        st.subheader("Tabla de Contenidos")
                        st.write("1. " + "\n2. ".join([sec['title'] for sec in table_of_contents]))
                    else:
                        st.error("No se pudo generar la tabla de contenidos.")
                else:
                    st.error("No se pudo generar el título y la descripción.")
else:
    st.info("El título, descripción y tabla de contenidos ya han sido generados. Si deseas generar nuevamente, reinicia la generación.")

# Permitir edición de la información inicial si ya se ha generado
if st.session_state.title and st.session_state.description and st.session_state.table_of_contents:
    st.markdown("---")
    st.header("Editar Información Inicial")
    
    with st.form("edit_initial_info"):
        # Editar Título del Estudio
        edited_title = st.text_input("Título del Estudio:", st.session_state.title)
        
        # Editar Descripción del Estudio
        edited_description = st.text_area("Descripción del Estudio:", st.session_state.description, height=200)
        
        # Editar Tabla de Contenidos
        st.subheader("Tabla de Contenidos")
        edited_table = []
        for sec in st.session_state.sections:
            edited_title_sec = st.text_input(f"Sección {sec['number']}:", sec['title'], key=f"sec_{sec['number']}")
            edited_table.append({"number": sec['number'], "title": edited_title_sec, "content": sec['content'], "references": sec['references']})
        
        submit_edit = st.form_submit_button("Guardar Cambios")
        
        if submit_edit:
            st.session_state.title = edited_title
            st.session_state.description = edited_description
            st.session_state.table_of_contents = edited_table
            st.session_state.sections = edited_table
            # Reconstruir el contenido Markdown
            st.session_state.markdown_content = f"# {st.session_state.title}\n\n{st.session_state.description}\n\n## Tabla de Contenidos\n\n"
            for sec in st.session_state.table_of_contents:
                st.session_state.markdown_content += f"{sec['number']}. {sec['title']}\n"
            st.session_state.markdown_content += "\n"
            st.success("Información inicial actualizada exitosamente.")

# Mostrar la sección para generar análisis solo si el título, descripción y tabla de contenidos han sido generados
if st.session_state.title and st.session_state.description and st.session_state.table_of_contents:
    st.markdown("---")
    st.header("Generación de Secciones de Análisis")

    # Barra de progreso
    generated_sections = len([sec for sec in st.session_state.sections if sec['content']])
    progress = generated_sections / st.session_state.total_sections
    progress_bar = st.progress(progress)

    # Botón para regenerar la sección seleccionada
    if st.session_state.selected_section is not None:
        with st.container():
            section = st.session_state.sections[st.session_state.selected_section]
            if section['title']:
                st.subheader(f"Sección {section['number']}: {section['title']}")
            else:
                st.subheader(f"Sección {section['number']}")
            if section['content']:
                # Mostrar contenido de la sección sin las referencias para evitar redundancia
                if "Referencias" in section['content']:
                    main_content, ref_content = section['content'].split("Referencias", 1)
                    st.markdown(main_content)
                    with st.expander("Ver Referencias"):
                        st.markdown(ref_content)
                else:
                    st.markdown(section['content'])
            else:
                st.info("Esta sección aún no ha sido generada.")
            if st.button("Regenerar Sección"):
                with st.spinner(f"Regenerando sección {section['number']}..."):
                    sec_num = section['number']
                    # Generar un nuevo título para la sección si es necesario
                    new_title = generate_section_title(work_title, work_type, sec_num)
                    if not new_title:
                        st.error(f"No se pudo generar el título para la sección {sec_num}.")
                    else:
                        # Generar contenido para la sección
                        new_content = generate_section(work_title, work_type, sec_num)
                        if new_content:
                            # Extraer referencias del nuevo contenido
                            new_references = extract_references(new_content)
                            # Actualizar la sección con el nuevo título, contenido y referencias
                            st.session_state.sections[st.session_state.selected_section]['title'] = new_title
                            st.session_state.sections[st.session_state.selected_section]['content'] = new_content
                            st.session_state.sections[st.session_state.selected_section]['references'] = new_references
                            
                            # Agregar referencias al estado global de referencias si no están duplicadas
                            for ref in new_references:
                                if ref not in st.session_state.references:
                                    st.session_state.references.append(ref)
                            
                            # Actualizar el contenido Markdown
                            if "Referencias" in section['content']:
                                chapter_heading_old = f"## Sección {sec_num}: " + (section['title'] if section['title'] else f"Sección {sec_num}") + "\n\n"
                                chapter_main_old, _ = section['content'].split("Referencias", 1)
                            else:
                                chapter_heading_old = f"## Sección {sec_num}: " + (section['title'] if section['title'] else f"Sección {sec_num}") + "\n\n"
                                chapter_main_old = section['content']
                            
                            chapter_heading_new = f"## Sección {sec_num}: {new_title}\n\n"
                            chapter_content_new = new_content + "\n\n"
                            
                            # Reemplazar en markdown_content
                            st.session_state.markdown_content = st.session_state.markdown_content.replace(
                                chapter_heading_old + chapter_main_old,
                                chapter_heading_new + chapter_content_new
                            )
                            
                            st.success(f"Sección {sec_num} regenerada exitosamente.")
                            # Actualizar la barra de progreso si la sección antes no tenía contenido
                            if not section['content']:
                                generated_sections += 1
                                progress = generated_sections / st.session_state.total_sections
                                progress_bar.progress(progress)
                        else:
                            st.error(f"No se pudo generar el contenido para la sección {sec_num}.")

    # Botón para generar la siguiente sección
    if st.session_state.current_section <= st.session_state.total_sections:
        if st.button("Generar Siguiente Sección"):
            with st.spinner(f"Generando sección {st.session_state.current_section}..."):
                section = st.session_state.sections[st.session_state.current_section - 1]
                if not section['content']:
                    # Generar título para la sección
                    generated_title = generate_section_title(work_title, work_type, st.session_state.current_section)
                    if not generated_title:
                        st.error(f"No se pudo generar el título para la sección {st.session_state.current_section}.")
                    else:
                        # Generar contenido para la sección
                        generated_content = generate_section(work_title, work_type, st.session_state.current_section)
                        if generated_content:
                            # Extraer referencias del contenido generado
                            generated_references = extract_references(generated_content)
                            # Actualizar la sección con título, contenido y referencias
                            st.session_state.sections[st.session_state.current_section - 1]['title'] = generated_title
                            st.session_state.sections[st.session_state.current_section - 1]['content'] = generated_content
                            st.session_state.sections[st.session_state.current_section - 1]['references'] = generated_references
                            
                            # Agregar referencias al estado global de referencias si no están duplicadas
                            for ref in generated_references:
                                if ref not in st.session_state.references:
                                    st.session_state.references.append(ref)
                            
                            # Agregar al contenido Markdown
                            section_heading = f"## Sección {st.session_state.current_section}: {generated_title}\n\n"
                            section_content = generated_content + "\n\n"
                            st.session_state.markdown_content += f"{section_heading}{section_content}"
                            
                            st.success(f"Sección {st.session_state.current_section} generada exitosamente.")
                            st.session_state.current_section += 1
                            
                            # Actualizar la barra de progreso
                            generated_sections += 1
                            progress = generated_sections / st.session_state.total_sections
                            progress_bar.progress(progress)
                        else:
                            st.error(f"No se pudo generar el contenido para la sección {st.session_state.current_section}.")
                else:
                    st.info(f"La sección {st.session_state.current_section} ya ha sido generada.")

    # Actualizar la barra de progreso hasta completar
    if len([sec for sec in st.session_state.sections if sec['content']]) == st.session_state.total_sections:
        st.session_state.generation_complete = True

    # Botones para exportar a Word
    if st.session_state.sections:
        with st.spinner("Preparando la exportación..."):
            # Exportar contenido parcial
            partial_markdown = f"# {st.session_state.title}\n\n{st.session_state.description}\n\n## Tabla de Contenidos\n\n"
            for sec in st.session_state.sections:
                if sec['content']:
                    partial_markdown += f"{sec['number']}. {sec['title']}\n"
            partial_markdown += "\n"
            for sec in st.session_state.sections:
                if sec['content']:
                    partial_markdown += f"## Sección {sec['number']}: {sec['title']}\n\n{sec['content']}\n\n"
                else:
                    break  # Solo incluir secciones generadas hasta el momento
            # Agregar referencias al final si hay
            if st.session_state.references:
                partial_markdown += "## Referencias\n\n"
                for ref in st.session_state.references:
                    partial_markdown += f"{ref}\n"
            partial_word_file = export_to_word(partial_markdown, st.session_state.references)
            
            # Exportar contenido completo
            if st.session_state.generation_complete:
                complete_markdown = st.session_state.markdown_content
                # Agregar todas las referencias al final
                complete_markdown += "\n## Referencias\n\n"
                for ref in st.session_state.references:
                    complete_markdown += f"{ref}\n"
                complete_word_file = export_to_word(complete_markdown, st.session_state.references)
                st.success("Preparado para descargar el estudio completo.")
            else:
                complete_word_file = None

            # Botón para descargar contenido parcial en Word
            st.download_button(
                label="Descargar Contenido Parcial en Word",
                data=partial_word_file,
                file_name=f"{st.session_state.title.replace(' ', '_')}_Parcial.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Botón para descargar contenido completo en Word (solo si está completo)
            if st.session_state.generation_complete:
                st.download_button(
                    label="Descargar Estudio Completo en Word",
                    data=complete_word_file,
                    file_name=f"{st.session_state.title.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # Opcional: Mostrar todo el contenido generado
    with st.expander("Mostrar Contenido Completo"):
        st.markdown(st.session_state.markdown_content)
